#Импорт библиотек
import docx
import matplotlib.pyplot as mat
#Создание переменных
d = docx.Document('lion.docx')
t = []
al = ''
k = {}
w = {}

#Заполнение t текстом по параграфам
for i in d.paragraphs:
    t.append(i.text)  
t = str(t)
#Удаление лишних символов
for i in range(33,64):
    t=t.replace(chr(i),'')
t=t.lower()                      
t = [t]
#разделение слов
for i in t:                 
    al += i
t = al.split()
#Работа со словарём из слов
for i in t:
    if i in k:
        k[i] += 1
    else:
        k[i] = 1
#работа со словарём из букв
for i in k:
    for j in i:
        if j in w:
            w[j] += 1
        else:
            w[j] = 1
#Параметры файла и таблицы
d = docx.Document()
tab = d.add_table(rows = 3,cols = 3)
tab.style = 'Table Grid'
#заполнение таблицы
tab.cell(0, 0).text = 'Слово'
tab.cell(0, 1).text = 'Частота встречи, раз'
tab.cell(0, 2).text = 'Частота встречи в %'
tab.cell(1, 0).text = str('солнце')
tab.cell(1, 1).text = str(k['солнце'])
tab.cell(1, 2).text = str(k['солнце']/len(t)*100)
tab.cell(2, 0).text = str('солнцу')
tab.cell(2, 1).text = str(k['солнцу'])
tab.cell(2, 2).text = str(k['солнцу']/len(t)*100)
#сохранение файла
d.save('sun.docx')

#Присвоение ключей и переменных для графика
b = list(w.keys())
a = list(w.values())
#Создание графика
mat.bar(b, a)
mat.xlabel('Буква')
mat.ylabel('Частота')
#Вывод графика
mat.show()
